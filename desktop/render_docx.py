# -*- coding: utf-8 -*-
"""Відтворення паперового бланка у форматі Word.

Функції set_cell / band / block і ширини колонок перенесено з генератора,
яким зроблено чинні плани — розмітка збігається з бланком установи.
"""
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

WIDTHS = [Cm(1.9), Cm(2.7), Cm(9.4), Cm(2.7), Cm(2.3)]
FONT = "Times New Roman"

_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "just": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _set_cell(cell, text, bold=False, size=10, align="just"):
    cell.text = ""
    par = cell.paragraphs[0]
    par.paragraph_format.space_after = Pt(0)
    par.alignment = _ALIGN[align]
    run = par.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = FONT


def build(doc_model, path):
    tpl = doc_model.tpl
    d = Document()

    st = d.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.space_after = Pt(0)
    st.paragraph_format.space_before = Pt(0)
    st.paragraph_format.line_spacing = 1.0

    sec = d.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin, sec.right_margin = Cm(1.5), Cm(1.0)
    sec.top_margin, sec.bottom_margin = Cm(1.5), Cm(1.5)

    def para(text="", bold=False, size=11, align="left", space_after=0):
        p = d.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.alignment = _ALIGN[align]
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        return p

    # ---------------- шапка ----------------
    para(tpl["doc_title"], bold=True, size=13, align="center", space_after=8)

    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Засудженого ")
    r.bold = True
    r.font.size = Pt(11)
    r = p.add_run(doc_model.title_line())
    r.font.size = Pt(11)
    r.underline = True
    para("(прізвище, власне ім'я, по батькові (за наявності))", size=8,
         align="center", space_after=2)

    date_from, date_to = doc_model.period()
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, underline, bold in (("на період з ", False, True),
                                  (date_from or "______________", True, False),
                                  ("  по  ", False, True),
                                  (date_to or "______________", True, False)):
        r = p.add_run(text)
        r.bold = bold
        r.underline = underline
        r.font.size = Pt(11)

    need_label = (tpl.get("need_field") or {}).get("label", "Криміногенна потреба")
    t0 = d.add_table(rows=2, cols=1)
    t0.style = "Table Grid"
    t0.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell(t0.rows[0].cells[0], need_label, bold=True, size=11, align="center")
    _set_cell(t0.rows[1].cells[0], doc_model.need.strip(), size=10)
    d.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---------------- основна таблиця ----------------
    table = d.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    header = ["Мета/цілі", "Цілі",
              "Поступові заходи, здійснення яких дадуть змогу зменшити/усунути "
              "актуальні фактори ризику",
              "Термін виконання", "Отриманий результат (виконано, внесено зміни)"]
    for i, txt in enumerate(header):
        _set_cell(table.rows[0].cells[i], txt, bold=True, size=10, align="center")

    def band(text):
        row = table.add_row()
        _set_cell(row.cells[0].merge(row.cells[4]), text, bold=True, size=10, align="left")

    def block(goals, term, obstacles):
        r1 = table.add_row().cells
        _set_cell(r1[0], "")
        _set_cell(r1[1], "Проміжні цілі", size=10, align="left")
        _set_cell(r1[2], goals, size=10)
        _set_cell(r1[3], term, size=10, align="center")
        _set_cell(r1[4], "")
        for c in table.add_row().cells:
            _set_cell(c, "")
        r3 = table.add_row().cells
        _set_cell(r3[0], "")
        _set_cell(r3[1], "Перепони та їх можливе вирішення", size=10, align="left")
        _set_cell(r3[2], obstacles, size=10)
        _set_cell(r3[3], "", size=10, align="center")
        _set_cell(r3[4], "")
        for c in table.add_row().cells:
            _set_cell(c, "")

    groups = [("during", "Наміри та плани засудженого під час відбування кримінального покарання"),
              ("after", "Наміри та плани засудженого після звільнення")]
    for group, caption in groups:
        specs = [s for s in tpl["sections"] if s["group"] == group]
        if not specs:
            continue
        band(caption)
        for spec in specs:
            data = doc_model.sections[spec["id"]]
            if data["skipped"] and not data["goals"].strip():
                continue
            band(f"{spec['number']}. {spec['title']}")
            block(data["goals"].strip(), data["term"].strip(), data["obstacles"].strip())

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = WIDTHS[i]

    # ---------------- підписи ----------------
    d.add_paragraph().paragraph_format.space_after = Pt(2)
    line = "______________________________________________   ____ ____________ 20___ р."
    for caption in ("(підпис, власне ім'я та прізвище начальника відділення СПС)",
                    "(підпис, власне ім'я та прізвище засудженого)"):
        p = d.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.add_run("План розробив " + line).font.size = Pt(11)
        para("                             " + caption, size=8, space_after=6)

    t2 = d.add_table(rows=2, cols=1)
    t2.style = "Table Grid"
    _set_cell(t2.rows[0].cells[0],
              "Висновки щодо результатів реалізації індивідуального плану\n"
              "виправлення та ресоціалізації", bold=True, size=11, align="center")
    _set_cell(t2.rows[1].cells[0], "\n\n\n", size=10)

    d.add_paragraph().paragraph_format.space_after = Pt(2)
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run("Начальник відділення СПС ______________________________   "
              "____ ____________ 20___ р.").font.size = Pt(11)
    para("                                              (підпис, власне ім'я та прізвище)",
         size=8, space_after=6)
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run("Ознайомлений ___________________________________________   "
              "____ ____________ 20___ р.").font.size = Pt(11)
    para("                                       (підпис, власне ім'я та прізвище засудженого)",
         size=8)

    cp = d.core_properties
    cp.title = tpl["doc_title"]
    cp.author = doc_model.pib() or ""
    cp.last_modified_by = doc_model.pib() or ""

    d.save(path)
    return path
